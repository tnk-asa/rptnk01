#---PMDA新記載要領XMLをdocx形式に変換する
#---XMLファイルからXSLTでHTML形式に変換した上でDOCXに展開
#---2022.03.09作成 by Tanikawa
#---2022.06.16更新 by Tanikawa Run出力処理の一本化
#---2022.07.25更新 by Tanikawa XMLから付加情報を抽出しjsonに出力
#---2022.09.09更新 by Tanikawa 順序無しリスト対応 箇条項目に・を表示

import sys
import re
import os
import glob
import shutil
import subprocess
from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from tkinter import filedialog
import lxml.etree as ET
import bs4
from bs4 import BeautifulSoup
from bs4 import Tag
import docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor        # 単位系や色が定義されているSharedクラス
from docx.enum.dml import MSO_THEME_COLOR   # 各種プロパティの設定が定義されているEnumerationsクラス
from docx.enum.text import WD_UNDERLINE
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Mm
from docx.shared import Cm, Inches
from PIL import Image

#-----------------------------------------------------------------------------
# フォルダ設定 プログラムファイルのパス取得
def get_dir_path(relative_path):
    try:
        base_path = sys._MEIPASS
        print("[Base Path (get from sys)]" + base_path)
    except Exception:
        base_path = os.path.dirname(__file__)
        print("[Base Path (get from sys)]" + base_path)
    return os.path.join(base_path, relative_path)
#-----------------------------------------------------------------------------
# 作業フォルダ設定
global wrk_dir, edt_dir, out_dir

current_file = get_dir_path(sys.argv[0])
wrk_dir = os.path.dirname(os.path.abspath(current_file))
#wrk_dir = os.path.abspath(os.path.dirname(__file__)) # .pyファイル用
#wrk_dir = os.path.dirname(sys.executable)           # .exe用
edt_dir = wrk_dir + "\edit"     # 中間ファイル処理
out_dir = wrk_dir + "\output"   # 出力先
#フォルダチェック
if not os.path.exists(out_dir):
    os.makedirs(out_dir) # 出力フォルダが存在しない場合は作成する

#-----------------------------------------------------------------------------
# GUI用 ファイル指定
def fileselect_clicked():
    fTyp = [("XMLファイル","*.xml"),("すべてのファイル","*")]
    filepath = filedialog.askopenfilename(filetypes = fTyp,initialdir = wrk_dir)
    file_name.set(filepath)
#------------------------------------------------------------------------
# GUI用 フォルダ指定
def dirdialog_clicked():
    iDirPath = filedialog.askdirectory(initialdir = wrk_dir)
    sav_dir.set(iDirPath)
#------------------------------------------------------------------------
# GUI用 保存先フォルダを開く
def OpenFolder(savedirectory):
    opn_dir = re.sub('/', '\\\\', savedirectory)
    subprocess.run('explorer {}'.format(opn_dir))
#------------------------------------------------------------------------
# PMDA XMLのNS対応
def file2pmdaxml(inFile):
    pmdaXmlStr = b'<?xml version="1.0" encoding="utf-8"?><ns:PackIns xmlns:ns="http://info.pmda.go.jp/namespace/prescription_drugs/package_insert/1.0"  xmlns="http://info.pmda.go.jp/namespace/prescription_drugs/package_insert/1.0" version="1.0" xmlKind="Packins" drugType="" referenceOfPrecautionsForUse="" referenceOfPrecautionsForHandling=""></ns:PackIns>'

    inFileLines = []
    with open(inFile, encoding='utf-8') as f:
        inFileLines = f.readlines()

    inFileStr = ''
    for line in inFileLines:
        inFileStr += line.strip()

    inFileXml = ET.fromstring(inFileStr.encode('utf-8'))
    inFileXmlRoot = inFileXml

    pmdaXml = ET.fromstring(pmdaXmlStr)
    pmdaXmlRoot = pmdaXml
    pmdaXmlRoot.set('drugType', inFileXmlRoot.get('drugType', default=""))
    pmdaXmlRoot.set('referenceOfPrecautionsForUse', inFileXmlRoot.get('referenceOfPrecautionsForUse', default=""))
    pmdaXmlRoot.set('referenceOfPrecautionsForHandling', inFileXmlRoot.get('referenceOfPrecautionsForHandling', default=""))
    pmdaXmlRoot.extend(list(inFileXmlRoot))

    return pmdaXmlRoot
#------------------------------------------------------------------------
# 付加情報抽出
def PickoutAppendix(soup_x, appfilename):
    app_itemlist = ['PackIns',
                    'PackageInsertNo',                          # 添付文書番号
                    'CompanyIdentifier',                        # 企業コード
                    'PreparationOrRevision',                    # 作成または改訂年月
                    'SccjNo',                                   # 日本標準商品分類番号（複数のケースあり）
                    'TherapeuticClassification',                # 薬効分類名
                    'DetailBrandName',                          # 販売名ごとの詳細
                    'GenericName',                              # 一般名
                    'SupplementaryInformaitonOfVaccineStrain',  # ワクチン株の補足情報　
                    'SpeciallyDescribedItems'                   # 特殊記載項目　SpeciallyDescribedItems
                    ]
    app_out = open(appfilename, 'w', encoding='utf-8')
    app_out.write('{')
    for app_item in app_itemlist:
        app_out_str = ''
        packins_1 = soup_x.find(app_item)
        # 照合元データID、添付文書形式
        if app_item == 'PackIns':
            for ref_datas in (['referenceOfPrecautionsForUse', 'referenceOfPrecautionsForHandling','drugType']):
                if packins_1.get(ref_datas) != None:
                    refdata_id = packins_1.get(ref_datas)
                else:
                    refdata_id = ''
                app_out_str = '\n    "' + ref_datas + '":' + '"' + refdata_id + '",'
                app_out.write(app_out_str)
                app_out_str = ''
        # 作成または改訂年月
        elif app_item == "PreparationOrRevision":
            packins_4 = soup_x.find_all('PreparationOrRevision')
            for date_rev in packins_4:
                if date_rev.get('id') == "今回":
                    app_out.write('\n''    "PreparationOrRevision":{')
                elif date_rev.get('id') == "前回":
                    app_out.write('\n''    "PreparationOrRevisionLast":{')
                else:
                    continue
                rev_yyyymm = rev_ver = rev_reason = ''
                for rev_subitem in date_rev.find_all(['YearMonth','Version','ReasonForRevision']):
                    if rev_subitem.name == 'YearMonth':
                        rev_yyyymm = rev_subitem.get_text()
                    if rev_subitem.name == 'Version':
                        rev_ver = re.sub("[\r\n]+", "", rev_subitem.get_text())
                    if rev_subitem.name == 'ReasonForRevision':
                        rev_reason = re.sub("[\r\n]+", "", rev_subitem.get_text())
                app_out_str = '\n' + '        "YearMonth":"' + rev_yyyymm + '",' + \
                              '\n' + '        "Version":"' + rev_ver + '",' + \
                              '\n' + '        "ReasonForRevision":"' + rev_reason + '",'
                if date_rev.get('created') == "1":
                    app_out_str += '\n        "created":true'
                else:
                    app_out_str += '\n        "created":false'
                app_out.write(app_out_str)
                app_out.write('\n    },')
        #日本標準商品分類番号（複数のケースあり）
        elif app_item == "SccjNo":
            sc_sn = 0
            if packins_1 != None:
                for sccj_itm in soup_x.find_all('SccjNo'):
                    if sccj_itm.get_text() != None and sccj_itm.get_text() != "" :
                        app_out_str = '\n    "SccjNo' + str(sc_sn) + '":"' + sccj_itm.get_text() + '",'
                        app_out.write(app_out_str)
                        sc_sn += 1
            else:
                app_out_str = '\n    "SccjNo0":"",'
                app_out.write(app_out_str)
        # 販売名ごとの詳細
        elif app_item == "DetailBrandName":
            brd_sn = 0
            packins_5 = soup_x.find_all('DetailBrandName')
            for date_rev in packins_5:
                app_out_str = '\n    "DetailBrandName' + str(brd_sn) + '":{'
                app_out.write(app_out_str)            
                brand_name = yj_code = brand_eng = brand_kana = ''
                for brand_subitem in date_rev.find_all(['ApprovalBrandName','YJCode','TrademarkInEnglish','BrandNameInHiragana']):
                    if brand_subitem.name == 'ApprovalBrandName':
                        brand_name = re.sub("[\r\n]+", "", brand_subitem.get_text())
                    if brand_subitem.name == 'YJCode':
                        yj_code = re.sub("[\r\n]+", "", brand_subitem.get_text())
                    if brand_subitem.name == 'TrademarkInEnglish':
                        brand_eng = re.sub("[\r\n]+", "", brand_subitem.get_text())
                    if brand_subitem.name == 'BrandNameInHiragana':
                        brand_kana = re.sub("[\r\n]+", "", brand_subitem.get_text())
                app_out_str = '\n' + '        "ApprovalBrandName":"' + brand_name + '",' + \
                              '\n' + '        "YJCode":"' + yj_code + '",' + \
                              '\n' + '        "TrademarkInEnglish":"' + brand_eng + '",' + \
                              '\n' + '        "BrandNameInHiragana":"' + brand_kana + '"'
                app_out.write(app_out_str)
                app_out.write('\n    },')
                brd_sn += 1
        else:
            if packins_1 != None:
                app_detail = packins_1.get_text()
                app_detail = re.sub("[\r\n]+", "", app_detail)
            else:
                app_detail = ""

            app_out_str = '\n    "' + app_item + '":' + '"' + app_detail + '"'
            if app_item != "SpeciallyDescribedItems":
                app_out_str += ','

            app_out.write(app_out_str)

    app_out.write('\n}')
    app_out.close()
#------------------------------------------------------------------------
# 改行・行末空白削除
def RemoveNewline(prm_str):
    cnv_txt = re.sub("^([\\r\\n])", "", prm_str)
    cnv_txt = re.sub("([\\r\\n])+", "", cnv_txt)
    cnv_txt = re.sub("^( |　|\\t)+$", "", cnv_txt)
    return cnv_txt
#------------------------------------------------------------------------
# header用セル出力
def Cell_output(tablename, tablerow, tablecol, celltext, fontname, fontsize):
    celltext = re.sub('\u00a0','\u0020', celltext)
    crn_cel = tablename.cell(tablerow, tablecol)
    cel_para = crn_cel.paragraphs
    cel_pgh = cel_para[0]
    cel_run = cel_pgh.add_run(celltext)
    cel_pgh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if fontname == '':
        cel_run.font.name = '游明朝'
    else:
        cel_run.font.name = fontname
    cel_run._element.rPr.rFonts.set(qn('w:eastAsia'), cel_run.font.name)
    fontsize = int(fontsize)
    if fontsize == 0 or fontsize == None:
        fontsize = 8
    cel_run.font.size = docx.shared.Pt(fontsize)
#------------------------------------------------------------------------
def HeaderRunOut(elements, itemname):
    if itemname == 'PreparationOrRevision': # 作成または改訂年月は文書の先頭に挿入
        try:
            hed_para = document.paragraphs[-1] # 最後の段落
        except:
            hed_para = document.add_paragraph()
    elif itemname == 'SpeciallyDescribedItems': # 特殊記載項目
        newline_para = document.add_paragraph()
        fram_obj = document.add_table(1, 1, "Table Grid")
        fram_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
        fcel_obj = fram_obj.cell(0,0)
        cel_para = fcel_obj.paragraphs
        hed_para = cel_para[0]
    else:
        hed_para = document.add_paragraph()

    # 項目別設定
    # 中央揃え
    if itemname in ['StandardName', 'TherapeuticClassification', 'GenericName', 'ApprovalBrandName', 'TrademarkName', 'SupplementaryInformationOfProduct', 'SupplementaryInformaitonOfVaccineStrain']:
        hed_para.alignment = WD_ALIGN_PARAGRAPH.CENTER    
    #フォント選択
    if itemname in ['ApprovalBrandName', 'TrademarkName']:
        run_font = "游ゴシック"
    else:    
        run_font = "游明朝" 
    #フォントサイズ
    if itemname in ['StandardName', 'TherapeuticClassification', 'GenericName', 'TrademarkName']:
        fnt_pt = 12
    elif itemname == 'ApprovalBrandName':
        fnt_pt = 16
    else:
        fnt_pt = 8
    # 文字装飾設定
    str_sup = str_sub = str_ita = str_bld = 0

    for elem in elements:
        elem_str = str(elem)
        if elem_str in [None, '', '\n', '\r']:
            continue # 空要素はスキップ
        if re.compile('enter').search(elem_str):
            run = hed_para.add_run("\n")
            continue # enter要素は改行のみ出力して次へ
        tag_jdg = re.search('<(.*?)>', elem_str) # タグ要素か否かのチェック
        if tag_jdg != None and elem_str != '' :
            if elem.name == 'Sup':
                str_sup = 1
            if elem.name == 'Sub':
                str_sub = 1
            if elem.name == 'Italic':
                str_ita = 1
            if elem.name == 'Bold':
                str_bld = 1
            if elem.name == 'Detail' and hed_para.text != '': #paragraphが空ではない状態でDtailタグが出現したら改行
                run = hed_para.add_run("\n")
        else:
        #RUN出力
            elem_str = re.sub('\u00a0','\u0020', elem_str) # nbspを通常空白へ置換
            run = hed_para.add_run(elem_str)
            run.font.name = run_font
            run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
            run.font.size = docx.shared.Pt(fnt_pt)
            run.font.color.rgb = RGBColor(0, 0, 0)
            if str_sup == 1:
                run.font.superscript = True
            else:
                run.font.superscript = False
            if str_sub == 1:
                run.font.subscript = True
            else:
                run.font.subscript = False
            if str_ita == 1:
                run.font.italic = True
            else:
                run.font.italic = False
            if str_bld == 1:
                run.font.bold = True
            else:
                run.font.bold = False      
            str_sup = str_sub = str_ita = str_bld = 0 

    if itemname == 'RegulatoryClassification':
        run.add_break(WD_BREAK.COLUMN)
#------------------------------------------------------------------------
# 出力文字列の編集
def Editstrings(prm_str):
    cnv_txt = re.sub("([\\r\\n])\\t+", "", prm_str)
    cnv_txt = re.sub("\\t", "", cnv_txt)
    cnv_txt = re.sub("^([\\r\\n])+$", "", cnv_txt)
    cnv_txt = re.sub("^[\\r\\n]+", "", cnv_txt)
    cnv_txt = re.sub("[:：]+：", "：", cnv_txt)
    cnv_txt = re.sub('\u00a0','\u0020', cnv_txt) #通常空白が'preview-include.xsl'でnbspに変換される対策 
    #cnv_txt = re.sub('\&lt;','<', cnv_txt)
    #cnv_txt = re.sub('\&gt;','>', cnv_txt)
    cnv_txt = re.sub('---NEWLINE---', "\\n", cnv_txt)
    return cnv_txt
#------------------------------------------------------------------------
# 空行除去
def Removenewline(prm_str):
    cnv_txt = re.sub("([\\r\\n])+", "", prm_str)
    return cnv_txt
#------------------------------------------------------------------------
# 画像 サイズ調節
def ImageSizeAdjust(img_path):
    image_file = Image.open(img_path)
    w, h = image_file.size
    size_decide = w/100
    if size_decide > 3.5:
        size_decide = 3.5
    return size_decide
#------------------------------------------------------------------------
# メイン処理 変換開始ボタンクリックで開始
def convbutton_clicked(targetfile):
    print(targetfile, "変換処理開始")
    xsltfile = edt_dir + '\\preview.xsl'  # XSLTファイルのパス
    src_dir = os.path.dirname(targetfile) # 入力ファイルがあるディレクトリ
    try:
        os.path.exists(xsltfile)
    except FileExistsError:
        exit("ERROR終了:XSLTファイルが見つかりません")
    # 旧ファイルが残っている場合は削除
    old_files = glob.glob(edt_dir + '\\*.*', recursive=True)
    for del_file in old_files:
        if re.search('(sourcedocument\.xml|temp\.html|\.(gif|GIF|png|PNG|jpg|JPG|jpeg|JPEG|docx?|DOCX?))', del_file):
            os.remove(del_file)
    shutil.copy(targetfile, edt_dir + '\\sourcedocument.xml')
    # XSLTによる変換処理
    dom = file2pmdaxml(targetfile)
    #dom = ET.parse(targetfile)
    xslt = ET.parse(xsltfile)
    transform = ET.XSLT(xslt)
    preview = transform(dom)
    infile = ET.tostring(preview, encoding='unicode', method='html', pretty_print=True)
    
    tmp_path = edt_dir + '\\temp.html' #中間ファイルhtml
    with open(tmp_path, 'w', encoding="utf-8") as outfile:
    #outfile = open(tmp_path, 'w', encoding="utf-8")
        outfile.write(infile)
    #入力フォルダ内にある画像ファイルを出力フォルダへコピー
    img_lst = glob.glob(src_dir + '\\*.*', recursive=True)
    for elm_img in img_lst:
        if re.search('\.(gif|GIF|png|PNG|jpg|JPG|jpeg|JPEG)', elm_img):
            try:
                shutil.copy(elm_img, edt_dir)
            except Exception as e:
                print(e)
    print("HTMLへの変換完了")
    # ============================================
    print("XMLヘッダの変換開始")
    in_file = targetfile
    #print(in_file)
    outfile_name = re.sub("(.*\/|\..*)", "", in_file) #最終出力docxファイルの名称

    #XML解析
    with open(in_file,encoding="utf_8") as xml_file:
        soup = bs4.BeautifulSoup(xml_file, "xml")
    #soup = bs4.BeautifulSoup(open(in_file,encoding="utf_8"), "xml")    
    
    #付加情報抽出
    appendix_file_name = out_dir + '\\' + outfile_name + ".json" # 出力ファイル名
    PickoutAppendix(soup, appendix_file_name)
        
    # documentオブジェクト作成
    global document
    wrn_ex = soup.find("Warnings") #警告の有無判定：赤帯入りのテンプレート使用
    doc_tpl = edt_dir + '\\template\\template_warning.docx'
    if wrn_ex != None:
        document = Document(doc_tpl)
    else:
        document = Document()

    # Stylesオブジェクト(Styleオブジェクトのコンテナ)を取得する
    styles = document.styles

    # デフォルトフォント
    run_font = "游明朝"
    font_pt = 8
    clr_red = clr_grn = clr_blu = 0

    #セクションパラメータ設定
    section0 = document.sections[0]
    sectPr0 = section0._sectPr
    ##段組み設定
    cols = sectPr0.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')  # コラム数
    cols.set(qn('w:sep'), '0')  # separate line: 0 or 1 区切り線
    cols.set(qn('w:equalWidth'), '1')  # equal width: 0 or 1
    ##ページ余白設定
    pg_margin0 = sectPr0.xpath('./w:pgMar')[0]
    pg_margin0.set(qn('w:top'), '720')
    pg_margin0.set(qn('w:right'), '720')
    pg_margin0.set(qn('w:bottom'), '720')
    pg_margin0.set(qn('w:left'), '720')
    pg_margin0.set(qn('w:header'), '227')
    pg_margin0.set(qn('w:footer'), '227')
    pg_margin0.set(qn('w:gutter'), '0')

    # 販売名-ApprovalBrandName 販売名の数をチェック
    brnname_lst = []
    for brn_nam in soup.find_all('ApprovalBrandName'):
        brnname_txt = RemoveNewline(brn_nam.get_text())
        brnname_lst.append(brnname_txt)
    brn_cnt = len(brnname_lst)

    # Header項目出力処理　ここから -------------------------------------------------
    # 作成又は改訂年月
    ver_run = tmp_run = rev_id = rev_ctd = now_txt = pre_txt = ''
    for item_rev in soup.find_all('PreparationOrRevision'):
        if item_rev.name == 'PreparationOrRevision':
            rev_id = item_rev.get('id') #今回または前回
            rev_ctd = item_rev.get('created') #作成または改訂
            if rev_id == "今回":
                tmp_run = "*"
            if soup.find_all(True, modified=re.compile(r"^前回")): # 前回改訂箇所が1つでもあれば改訂記号*を1つ追加
                tmp_run = tmp_run + "*"
        for elm_rev in item_rev.find_all(['YearMonth','Version','ReasonForRevision']):
            if elm_rev.name == 'YearMonth':
                rsn_run = ver_run = ''
                tmp_run = tmp_run + elm_rev.get_text()
                tmp_run = re.sub('(-|\/)','年',tmp_run)
                tmp_run = re.sub('$','月',tmp_run)
                if rev_ctd == '1':
                    tmp_run = tmp_run + '作成'
                else:
                    tmp_run = tmp_run + '改訂'
            if elm_rev.name == 'Version':
                ver_run = elm_rev.get_text()
                ver_run = re.sub("[\r\n]+","",ver_run)
                tmp_run = tmp_run + "（" + ver_run + "）"
            if elm_rev.name == 'ReasonForRevision':
                tmp_run = re.sub("）", "", tmp_run)
                rsn_run = elm_rev.get_text()
                rsn_run = re.sub("[\r\n]+","",rsn_run)
                tmp_run = tmp_run + "、" + rsn_run + "）"
        if rev_id == "今回":
            now_txt = tmp_run
            tmp_run = ''
        if rev_id == "前回":
            pre_txt = "\n" + tmp_run
            tmp_run = ''

    #run出力
    tmp_run = now_txt + pre_txt
    HeaderRunOut([tmp_run], 'PreparationOrRevision')

    #貯法使用期限Storage
    strg_run = ''
    strg_nam = soup.find('Storage')
    if strg_nam != None:
        for sub2_item in strg_nam.find_all(['StorageMethod','ShelfLife','OtherInformation']):
            if sub2_item.name == 'StorageMethod': # 貯法
                StorageMethod_txt = RemoveNewline(sub2_item.get_text())
                stg_run1 = '貯法：' + StorageMethod_txt
                strg_run = strg_run + '\n' + stg_run1
            if sub2_item.name == 'ShelfLife':
                ShelfLife_txt = RemoveNewline(sub2_item.get_text())
                stg_run2 = '有効期間：' + ShelfLife_txt
                strg_run = strg_run + '\n' + stg_run2
            if sub2_item.name == 'OtherInformation':
                oth_inf1 = sub2_item.find('Header')
                oth_hdr = oth_inf1.get_text()
                oth_inf2 = sub2_item.find('Detail')
                oth_dtl = oth_inf2.get_text()
                ShelfLife_txt = RemoveNewline(sub2_item.get_text())
                stg_run3 = oth_hdr + '：' + oth_dtl
                stg_run3 = re.sub("[\n\s]+", "", stg_run3)
                strg_run = strg_run + '\n' + stg_run3
#    strg_run = stg_run1 + '\n' + stg_run2
    HeaderRunOut([strg_run], 'Storage')

    #規制区分-RegulatoryClassificationCode
    tmp_run = ''
    regc_nam = soup.find('RegulatoryClassification')
    if regc_nam != None:
        for sub2_item in regc_nam.find_all('RegulatoryClassificationCode'):
            regcls_txt = RemoveNewline(sub2_item.get_text())
            if regcls_txt == '1':
                reg_run = '毒薬'
            if regcls_txt == '2':
                reg_run = '劇薬'
            if regcls_txt == '3':
                reg_run = '麻薬'
            if regcls_txt == '4':
                reg_run = '向精神薬(第一種)'
            if regcls_txt == '5':
                reg_run = '向精神薬(第二種)'
            if regcls_txt == '6':
                reg_run = '向精神薬(第三種)'
            if regcls_txt == '7':
                reg_run = '覚醒剤'
            if regcls_txt == '8':
                reg_run = '覚醒剤原料'
            if regcls_txt == '9':
                reg_run = '習慣性医薬品\n　注意―習慣性あり'
            if regcls_txt == '10':
                reg_run = '特例承認医薬品\n　注意―特例承認医薬品'
            if regcls_txt == '11':
                reg_run = '処方箋医薬品（医師）\n　注意―医師の処方箋により使用すること'
            if regcls_txt == '12':
                reg_run = '処方箋医薬品（医師等）\n　注意―医師等の処方箋により使用すること'
            if regcls_txt == '13':
                reg_run = '生物由来製品'
            if regcls_txt == '14':
                reg_run = '特定生物由来製品'
            if tmp_run == '':
                tmp_run = reg_run
            else:
                tmp_run = tmp_run + "\n" + reg_run
        HeaderRunOut([tmp_run], 'RegulatoryClassification')

    #テーブルブロックを追加
    # 作成又は改訂年月と日本標準商品分類番号
    global top_tbl
    top_tbl = document.add_table(2, 1, "Table Grid") #テーブルブロックを追加
    for cell in top_tbl.columns[0].cells: # 列幅設定
        cell.width = Inches(1.8)
        
    top_tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
    # 0-0 : 作成又は改訂年月
    # 0-1、1-1 : 余白
    # 0-2 : 日本標準商品分類番号　見出し
    # 1-2 : 日本標準商品分類番号

    # 日本標準商品分類番号
    Cell_output(top_tbl, 0, 0, "日本標準商品分類番号", run_font, font_pt)
    for sccj_no in soup.find_all('SccjNo'):
        sccj_txt = sccj_no.get_text()
        Cell_output(top_tbl, 1, 0, sccj_txt, run_font, font_pt)

    clm_brk = document.add_paragraph()

    section = document.add_section()
    section.start_type = WD_SECTION.CONTINUOUS # 現在の位置から新しいセクションを開始
    sectPr = section._sectPr
    ##段組み設定
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '1')  # コラム数
    cols.set(qn('w:sep'), '0')  # separate line: 0 or 1 区切り線
    cols.set(qn('w:equalWidth'), '1')  # equal width: 0 or 1

    #基準名
    std_para = document.add_paragraph() #テーブルブロックを追加
    std_nam = soup.find('StandardName')
    if std_nam != None:
        for sub2_item in std_nam.find_all(['StandardNameCategoryCode','StandardNameDetail']):
            if sub2_item.name == 'StandardNameCategoryCode': # 基準名
                cat_cod = sub2_item.get_text()
                if cat_cod == '1':
                    stc_run = '日本薬局方　'
                if cat_cod == '2':
                    stc_run = '生物学的製剤基準　'
                if cat_cod == '3':
                    stc_run = '放射性医薬品基準　'
                if cat_cod == '4':
                    stc_run = '血液型判定用抗体基準　'
            if sub2_item.name == 'StandardNameDetail': # 基準名詳細
                tmp_dtl = RemoveNewline(sub2_item.get_text())
                tmp_run = stc_run + tmp_dtl
                stc_run = tmp_dtl = ''
        HeaderRunOut([tmp_run], 'StandardName')

    #薬効分類
    tpt_nam = soup.find('TherapeuticClassification')
    if tpt_nam != None:
        thera_clas = tpt_nam.descendants
        HeaderRunOut(thera_clas, 'TherapeuticClassification')
    
    # 承認等
    brnname_lst = []
    trademark_lst = []
    supinfo_lst = []
    for app_subitem in soup.find_all('DetailBrandName'):
        brand_name = trade_mark = sup_info = ''
        for elm_subitem in app_subitem.find_all(['ApprovalBrandName', 'TrademarkName', 'SupplementaryInformationOfProduct']):
            tmp_run = elm_subitem.get_text()
            tmp_run = re.sub('[\r\n]+','',tmp_run)
            if elm_subitem.name == 'ApprovalBrandName':
                brand_name = tmp_run
        # 販売名英語　取得
            if elm_subitem.name == 'TrademarkName':
                trade_mark = tmp_run
        # 製品の補足情報　取得
            if elm_subitem.name == 'SupplementaryInformationOfProduct':
                sup_info = tmp_run
        brnname_lst.append(brand_name)
        trademark_lst.append(trade_mark)
        supinfo_lst.append(sup_info)

    # 販売名
    for brn_elm in brnname_lst:
        brn_idx = brnname_lst.index(brn_elm)
        print("■販売名",brn_idx, brn_elm)
        HeaderRunOut([brn_elm], 'ApprovalBrandName')
        # 販売名英語表記
        if trademark_lst != []:
            HeaderRunOut([trademark_lst[brn_idx]], 'TrademarkName')
        # 製品の補足情報
        if supinfo_lst != []:
            HeaderRunOut([supinfo_lst[brn_idx]], 'SupplementaryInformationOfProduct')

    # 一般名
    gen_get = soup.find('GenericName')
    if gen_get != None:
        gen_name = gen_get.descendants
        HeaderRunOut(gen_name, 'GenericName')

    # 承認番号表
    global app_tbl
    app_tbl = document.add_table(3, brn_cnt+1, "Table Grid") #テーブルブロックを追加
    app_tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
    top_cel = 0
    for cell in app_tbl.columns[0].cells: # 列幅設定
        if top_cel == 0:
            cell.width = Inches(0.6)
        else:
            cell.width = Inches(0.7)
        top_cel += 1
        
    app_cls = ''
    for app_or_lcs in soup.find_all(['ApprovalNo', 'LicenseNo']):
        app_cls = app_or_lcs.name
    if app_cls == None or app_cls =='':
        app_cls = 'ApprovalNo'
    for app_cel in range(3):
        crn_cel = app_tbl.cell(app_cel, 0)
        if app_cel == 1:
            if app_cls == 'LicenseNo':
                cel_txt = "許可番号"
            else:
                cel_txt = "承認番号"
        elif app_cel == 2:
            cel_txt = "販売開始"
        else:
            cel_txt = ""
        Cell_output(app_tbl, app_cel, 0, cel_txt, run_font, font_pt)

    # 販売名ごと
    brnname_lst = []
    crn_col = 0 # 承認番号表 列カウント
    for app_subitem in soup.find_all('DetailBrandName'):
        crn_col += 1
        for elm_subitem in app_subitem.find_all(['ApprovalBrandName', 'ApprovalNo', 'StartingDateOfMarketing']):
            tmp_run = elm_subitem.get_text()
            tmp_run = re.sub('[\r\n]+','',tmp_run)
        # 承認番号表出力
            if elm_subitem.name == 'StartingDateOfMarketing':
                tmp_run = re.sub('(\-|\/)','年',tmp_run)
                tmp_run = re.sub('$','月',tmp_run)
                crn_row = 2
            if elm_subitem.name == 'ApprovalNo':
                crn_row = 1
            if elm_subitem.name == 'ApprovalBrandName':
                crn_row = 0
                brnname_lst.append(tmp_run)
            if elm_subitem.name == 'StartingDateOfMarketing' or elm_subitem.name == 'ApprovalNo' or elm_subitem.name == 'ApprovalBrandName':
                Cell_output(app_tbl, crn_row, crn_col, tmp_run, run_font, font_pt)
    #特殊記載項目
    spdesc_get = soup.find('SpeciallyDescribedItems')
    if spdesc_get != None:
        spdesc_name = spdesc_get.descendants
        HeaderRunOut(spdesc_name, 'SpeciallyDescribedItems')
    #特殊記載項目
    vac_info = soup.find('SupplementaryInformaitonOfVaccineStrain')
    if vac_info != None:
        vaccine_name = vac_info.descendants
        HeaderRunOut(vaccine_name, 'SupplementaryInformaitonOfVaccineStrain')

    #ヘッダと本文の間に区切り線
    div_para = document.add_paragraph()  
    div_img = edt_dir + '\\template\\dividing_line.png'
    div_para.add_run().add_picture(div_img, width=Mm(193), height=Mm(0.5))
#------------------------------------------------------------------------
# メイン項目出力
    #tst_dir = wrk_dir + '\\' + 'edit\\'
    in_file = edt_dir + '\\temp.html'
    #out_dir2 =  wrk_dir + '\\' + 'output'
    word_write_file_name = out_dir + '\\' + outfile_name + ".docx" # 出力ファイル名

    # ドキュメントオブジェクト作成
    styles = document.styles

    #セクションパラメータ設定
    section = document.add_section()
    section.start_type = WD_SECTION.CONTINUOUS # 現在の位置から新しいセクションを開始
    sectPr = section._sectPr
    ##段組み設定
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')  # コラム数
    cols.set(qn('w:sep'), '0')  # separate line: 0 or 1 区切り線
    cols.set(qn('w:equalWidth'), '1')  # equal width: 0 or 1

    #HTML解析
    with open(in_file,encoding="utf_8") as html_file:
        soup = bs4.BeautifulSoup(html_file, "html.parser")
        
    #組成・性状表補正・・・販売名の行を追加する
    for comp_and_prop in soup.find_all('div', attrs={'id': "HDR_CompositionAndProperty"}):
        for brdname_hed in comp_and_prop.find_all('div', attrs={'data-level': "3"}):
            brdname_hed2 = brdname_hed.find('h3')
            brdname_hed2.attrs['nobreak'] = 'yes' #セル先頭で要素を改行させないための属性追加
            for candp_table in brdname_hed.find_all('table', attrs={'class': "CompositionAndProperty_table"}):
                first_row = candp_table.find('tr')
                brd_row = BeautifulSoup(str(first_row), features="lxml") #1行目を複製して先頭行に追加
                # 複製した行を補正して追加
                col_num = len(brd_row.find_all(['th','td']))
                col_cnt = 0
                for row_dtl in brd_row.find_all(['th','td']):
                    col_cnt += 1
                    try:
                        del row_dtl.attrs['rowspan']
                    except:
                        pass
                    row_dtl.clear()
                    if col_cnt == 1:
                        brd_celhedr = soup.new_tag('h3', attrs={"class":"section_header", "nobreak":"yes"})
                        brd_celhedr.string = "販売名"
                        row_dtl.insert(1, brd_celhedr)
                        if col_num > 2:
                            row_dtl.attrs['colspan'] = '2'
                    elif col_cnt == 2 and col_num > 2:
                        row_dtl.decompose()
                    elif col_cnt == 3:
                        brow_clspn = col_num - 2
                        brow_clspns = str(brow_clspn)
                        row_dtl.attrs['colspan'] = brow_clspns
                        row_dtl.insert(1, brdname_hed2)
                    elif col_cnt > 3:
                        row_dtl.decompose()
                    else:
                        row_dtl.insert(1, brdname_hed2)
                first_row.insert_before(brd_row) # 先頭行の直前に追加
    #順序無しリスト　箇条書き変換 20220909
    for list_items in soup.find_all('ul', attrs={'style': "list-style-type: "}):
        for elm_lst_itm in list_items.find_all('li'):
            elm_lst_itm.insert(0, "・")

    #改訂記号を順序番号の前へ移動
    for rev_item in soup.find_all(['li', 'p', 'h3']):
        num_flg = 0
        if rev_item.name == 'p':
            if  rev_item.get('class') == ['wordBreak']:
                num_flg = 1
        for rev_mov in rev_item.find_all('span'):     # 改訂記号を順序番号の前へ出す 20220609
            if rev_mov.get('class') == ['section_header']:
                rev_mov.attrs['nobreak'] = 'yes' #改訂記号直後の要素を改行させないための属性追加 20220623
                num_flg = 1
                continue
            if rev_mov.get('class') in ( ['revisionThis-editor'] , ['revisionPrev-editor'] , ['revisionPrevThis-editor'] ):
                #rev_mov が改訂記号　まるごとタグ要素
                if rev_mov != None and num_flg > 0:
                    rev_item.insert(0, rev_mov)
#                else:
#                    rev_mov.unwrap()
            #ヘッダ要素の入れ子改訂記号をタグの外に出す
            if rev_mov.get('class') == ['Header-preview'] or rev_item.name == 'h3':
                if rev_item.name == 'h3':
                    revhed_txt = '---NEWLINE---**'
                    rev_child = rev_item.children
                    nst_revmv = rev_item
                else:
                    rev_child = rev_mov.children
                    revhed_txt = '**'
                    nst_revmv = rev_mov
                for revmark_nest in rev_child:
                    tag_type = str(type(revmark_nest))
                    if re.compile('bs4\.element\.Tag').search(tag_type):
                        if revmark_nest.get('class') in ( ['revisionThis-editor'] , ['revisionPrev-editor'] , ['revisionPrevThis-editor'] ):
                            nst_revmv.attrs['nobreak'] = 'yes' #改訂記号直後の要素を改行させないための属性追加 20220623
                            revmark_nest.clear()
                            revmark_nest.string = revhed_txt
                            mv_nstrev = revmark_nest.extract()
                            nst_revmv.insert_before(mv_nstrev)

    # 改訂記号が連続になった場合は先頭のもの以外を削除する 20220627
    for rev_srs in reversed(soup.find_all('span')): #リスト化したspan要素を逆順で読み込み
        if rev_srs.get('class') in ( ['revisionThis-editor'] , ['revisionPrev-editor'] , ['revisionPrevThis-editor'] ):
            sub_hdr = rev_srs.find_next_sibling()
            if sub_hdr != None:
                if sub_hdr.get('class') in ( ['revisionThis-editor'] , ['revisionPrev-editor'] , ['revisionPrevThis-editor'] ):
                    sub_hdr.decompose()

    ## 見出し参照変換
    for hdr_ref in soup.find_all('a', attrs={'class': "HeaderRef"}):
        hdr_num = ''
        str_ref = hdr_ref.get('href')
        str_ref = re.sub('#', '', str_ref)
        for hdr_nums in soup.find_all('div', attrs={'data-header-id': str_ref}):
            hdr_num = hdr_nums.get_text()
        ref_cnv = "[" + hdr_num + " 参照]"
        hdr_ref.append(ref_cnv)
        hdr_ref.unwrap()
    #文献参照　前処理
    for bok_ref in soup.find_all('sup', attrs={'class': "ReferenceBookRef"}):
        #print(bok_ref)
        ref_txt = bok_ref.get_text()
        ref_txt = re.sub("( |　|\t|[\r\n])+","",ref_txt)
        bok_ref.clear()
        bok_ref.append(ref_txt) 
    #コメント関連前処理
    for cmt_num in soup.find_all('span', attrs={'class': "CommentNum"}):
        cmt_num.unwrap()
    for cmt_num in soup.find_all('span', attrs={'class': "CommentRefNum"}):
        cmt_num.unwrap()
    for cmt_num in soup.find_all('sup', attrs={'class': "CommentRef"}):
        sup_txt = cmt_num.get_text()
        sup_txt = re.sub("( |　|\t|[\r\n])+","",sup_txt)
        cmt_num.clear()
        cmt_num.append(sup_txt)

    for lst_num in soup.find_all('span', attrs={'class': "section_header"}):
        sub_hdr = lst_num.find_next_sibling()
        if sub_hdr is not None and sub_hdr.get('class') == ['Header-preview']:
            sub_hdr.attrs['nobreak'] = 'yes'

    #画像タグの前処理
    for img_elm in soup.find_all('img'):
        img_name = img_elm.get('src')
        replace_img = img_elm.replace_with("ImageInsert:" + img_name)
        #print(replace_img)
    #強制改行処理
    for brake_tag in soup.find_all('br'):
        brake_tag.replace_with('---NEWLINE---')
    #空白要素削除
    for text_element in soup.find_all(string=re.compile(r'^\s+$')):
        text_element.replace_with('')
    for text_element in soup.find_all(['div', 'h3']):
        if text_element.get_text() == '' or text_element.get_text() == None:
            text_element.decompose()
    #空の文字装飾要素を削除--2022.06.20 空タグに対する対応
    for supsub_elem in soup.find_all(['sup','sub', 'em', 'i']):
        supsub_txt = supsub_elem.get_text()
        if re.compile('^(| |\\s)+$').search(supsub_txt):
            supsub_elem.decompose()

    #テーブル ブランク行 除去  # cell spanの関係でエラーの原因になる事がある 2022.06.01
    for blk_row in soup.find_all('tr'):
        row_txt = blk_row.get_text()
        if row_txt == '' or row_txt == None:
            blk_row.decompose()
    #テーブル終了位置
    for warn_end in soup.find_all('div', attrs={'id': 'HDR_Warnings'}):
        tag_warn = soup.new_tag('p')
        tag_warn.string = '---TABLEEND---'
        warn_end.insert_after(tag_warn)
    for tbl_elm in soup.find_all('table'):
        tag_tbl = soup.new_tag('p')
        tag_tbl.string = '---TABLEEND---'
        tbl_elm.insert_after(tag_tbl)

    item_list = [
     'HDR_Warnings',        # 01.警告 // //
     'HDR_ContraIndications',            # 02.禁忌 //
     'HDR_CompositionAndProperty',       # 03.組成・性状 //
     'HDR_IndicationsOrEfficacy',        # 04.効能又は効果 //
     'HDR_EfficacyRelatedPrecautions',   # 05.効能又は効果に関連する注意 //
     'HDR_InfoDoseAdmin',                # 06.用法及び用量 //
     'HDR_InfoPrecautionsDosage',        # 07.用法及び用量に関連する注意 //
     'HDR_ImportantPrecautions',         # 08.重要な基本的注意 //
     'HDR_UseInSpecificPopulations',     # 09.特定の背景を有する患者に関する注意 //
     'HDR_Interactions',                 # 10.相互作用 //
     'HDR_AdverseEvents',                # 11.副作用 //
     'HDR_InfluenceOnLaboratoryValues',  # 12.臨床検査結果に及ぼす影響 //
     'HDR_OverDosage',                   # 13.過量投与 //
     'HDR_PrecautionsForApplication',    # 14.適用上の注意 //
     'HDR_OtherPrecautions',             # 15.その他の注意 //
     'HDR_Pharmacokinetics',             # 16.薬物動態 //
     'HDR_ResultsOfClinicalTrials',      # 17.臨床成績 //
     'HDR_EfficacyPharmacology',         # 18.薬効薬理 //
     'HDR_PhyschemOfActIngredients',     # 19.有効成分に関する理化学的知見 //
     'HDR_PrecautionsForHandling',       # 20.取扱い上の注意 //
     'HDR_ConditionsOfApproval',         # 21.承認条件 //
     'HDR_Package',                      # 22.包装 //
     'HDR_MainLiterature',               # 23.主要文献 //
     'HDR_AddresseeOfLiteratureRequest', # 24.文献請求先及び問い合わせ先 //
     'HDR_AttentionOfInsurance',         # 25.保険給付上の注意 //
     'HDR_NameAddressManufact',          # 26.製造販売業者等 //
     'HDR_ReferenceInformation'          # 参考情報 //
    ]

    # 変数初期値
    itm_txt = ''
    tag_get_li = []
    str_dcr = ''
    fnt_bld = ''
    run_font = "游明朝"
    flg_newlin = 0 #改行フラグ
    fln_nonewlin = 0 #改行しないフラグ
    flg_newpara = 0 #新パラグラフフラグ
    clr_red = clr_grn = clr_blu = 0  #RGBカラー R,G,B値

    # メイン処理　-大項目ごとに編集
    for item_name in item_list:   #item_list:大項目名のリスト
        sub_items = []
        sub_items.clear()
        itm_lv1 = soup.find('div', attrs={'id': item_name})
        if itm_lv1 != None:
            print("-----",item_name,"-----")
            sub_items = itm_lv1.descendants #子孫要素をすべて取得するしくみ リストではなくジェネレータオブジェクトになる
        else:
            continue
        # 1.警告、2.禁忌の項は赤枠用に1行1列のテーブルオブジェクトを配置
        if item_name == "HDR_Warnings" or item_name == "HDR_ContraIndications":
            fram_obj = document.add_table(1, 1)
            #fram_obj.style = 'TableGrid'
            fcel_obj = fram_obj.cell(0,0)
            #罫線の色設定
            tbl_bdr = fcel_obj._tc
            tcl_prm = tbl_bdr.get_or_add_tcPr()
            bdr_atr = OxmlElement('w:tcBorders')
            tcl_prm.append(bdr_atr)
            for bdr_edg in ['top', 'left', 'right', 'bottom']:
                tag_elm = 'w:' + bdr_edg
                apd_tag = tag_elm.format(bdr_edg)
                elm_atr = bdr_atr.find(qn(apd_tag))
                if elm_atr is None:
                    elm_atr = OxmlElement(apd_tag)
                    bdr_atr.append(elm_atr)
                    elm_atr.set(qn('w:color'), "FF0000")
                    elm_atr.set(qn('w:val'), "single")
            
            cel_para = fcel_obj.paragraphs
            cmn_para = cel_para[0]
            fln_nonewlin = 1
            flg_newpara = 0
        else:
            cmn_para = document.add_paragraph() #1,2項以外は通常パラグラフを配置
            fln_nonewlin = 1
            flg_newpara = 0
        # 1.警告の文字色設定
        if item_name == "HDR_Warnings":
            clr_red = 255  #RGBカラー R
            clr_grn = 0  #RGBカラー G
            clr_blu = 0  #RGBカラー B
        else:
            clr_red = clr_grn = clr_blu = 0  #RGBカラー RGB

        # 大項目以下 要素ごとの処理を開始
        for elm_sub_item in sub_items:
            # 各値をデフォルトに戻すのはRun出力後にする 2022.06.20
            #if elm_sub_item.name in ['p','li','span','div']:
            #    run_font = "游明朝"
            #    fnt_bld = ''
            #    str_dcr = ''

            # 親要素チェック
            item_parent = elm_sub_item.parent
            if item_parent.name == 'h3' and item_parent.get('class') == ['section_header'] or   \
                item_parent.name == 'span' and item_parent.get('class') == ['Header-preview'] :
                    run_font = "游ゴシック"
                    fnt_bld = "bold"
            # タグ要素か否かをチェックして処理開始
            elm_sub_str = str(elm_sub_item) #要素を文字列化
            tag_jdg = re.search('<[A-Za-z](.*?)(=|>)', elm_sub_str) 
            if tag_jdg != None and elm_sub_str != '' :
                #フラグがONの場合は新しいパラグラフを追加
                if elm_sub_item.name != 'table' and flg_newpara == 1:
                    cmn_para = document.add_paragraph()
                    flg_newpara = 0
                    fln_nonewlin = 1
                # 表組出力
                if elm_sub_item.name == 'table':
                    #print(elm_sub_item,"▲▲▲▲▲Tableが見つかった▲▲▲▲▲")
                    flg_newpara = 0 #表を追加する場合は新しいパラグラフをオフ
                    flg_footer = 0 #Footerフラグ
                    footer_row = 0 #Footer行保持用
                    run_font = "游明朝"
                    fnt_bld = str_dcr = ''
                    
                    if elm_sub_item.get('class') == ['ContraIndication_table']: #赤罫線設定オン
                        tbl_red = 1
                    else:
                        tbl_red = 0
                    #行と列の処理
                    out_rows = []
                    rows_count = cols_count = 0 #行数、列数カウント
                    row_spn = col_spn = 0  #セル結合チェック 行、列
                    colg_nocnt = 0 #colgroupを行数に含めないための処理
                    chk_cel_lst = [] #被結合セルの番号格納用：スキップするためのチェック用
                    for tbl_row in elm_sub_item.find_all(['tr','colgroup']):
                        out_cols = []
                        col_spn_count = 0 # colspanで行数カウント 2022.06.01
                        for tbl_cel in tbl_row.find_all(['th','td','col']):
                            out_cols.append(tbl_cel)
                            #--- colspanによる列数カウント 2022.06.01 ---
                            try:
                                span_int = tbl_cel['colspan']
                                if span_int != "" and span_int != "0" and span_int != "99":
                                    col_spn_count += int(span_int)
                            except:
                                pass                            
                        #print("列数：",len(out_cols))
                        out_rows.append(out_cols)
                        if len(out_cols) > cols_count:
                            cols_count = len(out_cols) #列数は最大の行の値を使用する（セル結合対策）
                        if col_spn_count > cols_count:
                            cols_count = col_spn_count #colspanの合計値の方が大きい場合は置き換え 2022.06.01
                        if tbl_row.name == 'tr':
                            rows_count += 1

                    #print('★表の行数、列数',rows_count,cols_count)
                    tbl_obj = document.add_table(rows_count, cols_count, "Table Grid") #テーブルブロックを追加
                    current_row = -1
                    current_col = 0
                    rows_count -= 1 #リストで処理するため最終行の値 -1
                    cols_count -= 1 #リストで処理するため最終列の値 -1
                    continue
                if elm_sub_item.name == 'tr' or elm_sub_item.name == 'th' or elm_sub_item.name == 'td':
                    if elm_sub_item.name == 'tr':
                        row_txt = elm_sub_item.get_text() # 行がブランクの場合に処理を分岐するため取得保持 2022.06.01
                    #print(elm_sub_item.name, "--", elm_sub_item.get_text())
                    if current_col > cols_count or elm_sub_item.name == 'tr':
                        current_row += 1
                        current_col = 0
                        if elm_sub_item.get('class') == ['tableFooter']: #table フッター処理
                            footer_row = current_row
                    #print("■現在のセル",current_row,current_col,"=====================")
                    for cel_elm in chk_cel_lst:
                        if current_row == cel_elm[0] and current_col == cel_elm[1]:
                            #print(cel_elm,cel_elm[0],cel_elm[1],"このセルは吸収された",current_col)
                            current_col += 1

                    #----- セル処理 結合・配置 -----
                    if elm_sub_item.name == 'th' or elm_sub_item.name == 'td':
                        # align 設定用 style属性の値を取得
                        style_val = ''
                        try:
                            style_val = elm_sub_item['style']
                            if elm_sub_item.name == 'th':
                                if re.compile('align').search(style_val):
                                    pass
                                else:
                                    style_val = style_val + 'vertical-align:middle;text-align:center;'
                        except:
                            if elm_sub_item.name == 'th':
                                style_val = 'vertical-align:middle;text-align:center;'

                        #--- rowspan処理 ---
                        try:
                            cell_att = elm_sub_item['rowspan']
                            if cell_att != "" and cell_att != "0" and cell_att != "1":
                                row_spn = int(cell_att) - 1
                                #print("・セル-縦結合",row_spn)
                        except:
                            pass
                        #--- colspan処理 ---
                        try:
                            cell_att = elm_sub_item['colspan']
                            if cell_att != "" and cell_att != "0" and cell_att != "1":
                                col_spn = int(cell_att) - 1
                                #print("・セル-横結合",col_spn)
                            if cell_att == "99": #colspan=99設定の回避 2022.06.13
                                col_spn = cols_count
                        except:
                            pass
                        if row_spn != 0 or col_spn != 0:
                            spn_cl1 = tbl_obj.cell(current_row, current_col)
                            if rows_count < current_row + row_spn: # 空白行削除によるセル結合不整合対策 20220601
                                spn_cl2 = tbl_obj.cell(rows_count, current_col + col_spn)
                            else:
                                spn_cl2 = tbl_obj.cell(current_row + row_spn, current_col + col_spn)
                            #print(current_row, current_col,"～",current_row + row_spn, current_col + col_spn,"を結合した")
                            unt_cel = spn_cl1.merge(spn_cl2)
                        #--- 被結合セルをリスト化：被結合セルをスキップするためのチェック用 ---
                        for row_chk in range(current_row, current_row + row_spn + 1):
                            for col_chk in range(current_col, current_col + col_spn + 1):
                                if row_chk != current_row or col_chk != current_col:
                                    chk_cel = [row_chk, col_chk, col_spn]
                                    chk_cel_lst.append(chk_cel)
                        #print("●行列処理開始", current_row,current_col)
                        if row_txt != '' and row_txt != None:
                            try:
                                cel_obj = tbl_obj.cell(current_row,current_col) # テーブルにセルオブジェクトを追加
                                cel_para = cel_obj.paragraphs
                                cmn_para = cel_para[0]
                                # vertical align 設定
                                if re.compile('vertical-align:middle').search(style_val):
                                    cmn_para.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                elif re.compile('vertical-align:bottom').search(style_val):
                                    cmn_para.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                                else:
                                    cmn_para.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                                # text align 設定
                                if re.compile('text-align:center').search(style_val):
                                    cmn_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                elif re.compile('text-align:right').search(style_val):
                                    cmn_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                else:
                                    cmn_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            except Exception as e:
                                print(item_name, current_row,current_col, "◎セル結合でエラー◎", row_txt, e)

                        if current_row != 0 and current_row == footer_row and current_col == 0:
                            flg_footer = 1
                        else:
                            flg_footer = 0
                        
                        # 表罫線の色--------------------------------------------
                        tbl_bdr = cel_obj._tc
                        tcl_prm = tbl_bdr.get_or_add_tcPr()
                        bdr_atr = OxmlElement('w:tcBorders')
                        tcl_prm.append(bdr_atr)
                        
                        for bdr_edg in ['top', 'left', 'right', 'bottom']:
                            tag_elm = 'w:' + bdr_edg
                            apd_tag = tag_elm.format(bdr_edg)

                            elm_atr = bdr_atr.find(qn(apd_tag))
                            if elm_atr is None:
                                elm_atr = OxmlElement(apd_tag)
                                bdr_atr.append(elm_atr)

                            if tbl_red == 1: #表罫線のカラー設定
                                elm_atr.set(qn('w:color'), "FF0000") # 赤
                            elif flg_footer > 0:
                                elm_atr.set(qn('w:color'), "FFFFFF") # 白（不可視化）
                            else:
                                elm_atr.set(qn('w:color'), "000000") # 黒
                            elm_atr.set(qn('w:val'), "single")
                        #print("このセルの内容：",elm_sub_item.get_text(),"\n")
                        current_col =  current_col + 1
                        row_spn = 0  #セル結合チェック 行
                        col_spn = 0  #セル結合チェック 列
                    continue
                #画像の挿入
                if elm_sub_item.name == 'figure': #画像ブロックとして存在する場合はパラグラフ追加
                    flg_newpara = 1
                    continue
                #見出し処理1
                if elm_sub_item.name == 'h3' and elm_sub_item.get('class') == ['section_header'] or   \
                    elm_sub_item.name == 'span' and elm_sub_item.get('class') == ['section_header'] :
                    if elm_sub_item.get('nobreak') == 'yes':
                        fln_nonewlin = 1
                        #input("改行しない")
                    else:
                        flg_newlin = 1
                    if item_name == 'HDR_PhyschemOfActIngredients' and item_parent.name == 'div' and item_parent.get('data-level') == '99':
                        elm_sub_item.string.insert_after("：")  #19項サブ見出しに：を追加
                    run_font = "游ゴシック"
                    fnt_bld = "bold"
                    continue
                #見出し処理2
                if elm_sub_item.name == 'span' and elm_sub_item.get('class') == ['Header-preview']:
                    #input(elm_sub_item.get('nobreak'))
                    if elm_sub_item.get('nobreak') == 'yes':
                        fln_nonewlin = 1
                        #input("改行しない")
                    else:
                        flg_newlin = 1
                    run_font = "游ゴシック"
                    fnt_bld = "bold"
                    continue

                #改行処理
                if fln_nonewlin != 1 and elm_sub_item.get_text() != '':
                    if elm_sub_item.name == 'p' or elm_sub_item.name == 'li' or elm_sub_item.name == 'caption':
                        if cmn_para.text != '': #paragraph先頭の場合は改行処理しない
                            cmn_para.add_run('\n')
                            fln_nonewlin = 1
                        continue
                    if elm_sub_item.name == 'div' and elm_sub_item.get('class') == ['Comment']:
                        if cmn_para.text != '': #paragraph先頭の場合は改行処理しない
                            cmn_para.add_run('\n')
                            fln_nonewlin = 1
                        continue
                    if elm_sub_item.name == 'div' and elm_sub_item.get('data-level') == '99':
                        cmn_para.add_run('\n')
                        fln_nonewlin = 1
                        continue
                    if elm_sub_item.name == 'div' and elm_sub_item.get('data-index') in ['1','2','3','4','5','6','7','8','9','10','11','12','13','14','15','16','17','18','19','20','21','22','23','24','25','26']:
                        cmn_para.add_run('\n')
                        fln_nonewlin = 1
                        continue
                    if elm_sub_item.name == 'div' and elm_sub_item.get('class') == ['level-3'] and item_name == "HDR_CompositionAndProperty":
                        cmn_para.add_run('\n')
                        fln_nonewlin = 1
                        continue

                #文字装飾処理
                if elm_sub_item.name == 'span' or elm_sub_item.name == 'figure' or \
                elm_sub_item.name == 'sub' or elm_sub_item.name == 'sup' or \
                elm_sub_item.name == 'em' or elm_sub_item.name == 'i' :
                    #print("★改行しない",elm_sub_item)
                    fln_nonewlin = 1
                if elm_sub_item.name == 'sup':
                    str_dcr = 'sup'
                    continue
                if elm_sub_item.name == 'sub':
                    str_dcr = 'sub'
                    continue
                if elm_sub_item.name == 'em' or elm_sub_item.name == 'i':
                    str_dcr = 'italic'
                    continue
            elif elm_sub_str == '':
                #print("\nｘｘｘ文字列が含まれない", tag_jdg, elm_sub_str)
                continue

            ##### RUN出力処理 ##### 文字列要素のみになったところでRUN出力
            else:
                #画像出力
                if re.search("^ImageInsert:", elm_sub_str):
                    elm_sub_str = re.sub("ImageInsert:", "", elm_sub_str)
                    img_path = edt_dir + "\\" + elm_sub_str
                    img_size = ImageSizeAdjust(img_path)
                    if item_name == 'HDR_CompositionAndProperty':
                        if img_size > 1.5:
                            img_size = 1.5
                        cmn_para.add_run().add_picture(img_path, width=Inches(img_size))
                    else:
                        cmn_para.add_run('\n').add_picture(img_path, width=Inches(img_size))                        
                #表の終了処理
                elif elm_sub_str == '---TABLEEND---':
                    cmn_para = document.add_paragraph()
                #文字列出力
                else:
                    if fln_nonewlin == 1:
                        #print(elm_sub_str)
                        elm_sub_str = Removenewline(elm_sub_str)
                    if flg_newlin > 0 and fln_nonewlin == 0:
                        cmn_para.add_run('\n') #▲▲▲組成性状の表　先頭セルに改行が入ることを確認20220629
                    
                    elm_sub_str = Editstrings(elm_sub_str) #出力する文字列の下処理
                    #run出力
                    run = cmn_para.add_run(elm_sub_str)
                    run.font.name = run_font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
                    run.font.size = docx.shared.Pt(8)
                    run.font.color.rgb = RGBColor(clr_red, clr_grn, clr_blu)
                    if str_dcr == 'sup':
                        run.font.superscript = True
                    else:
                        run.font.superscript = False
                    if str_dcr == 'sub':
                        run.font.subscript = True
                    else:
                        run.font.subscript = False
                    if str_dcr == 'italic':
                        run.font.italic = True
                    else:
                        run.font.italic = False
                    if fnt_bld == 'bold':
                        run.font.bold = True
                    else:
                        run.font.bold = False
                str_dcr = fnt_bld = ''
                run_font = "游明朝"
                flg_newlin = 0
                fln_nonewlin = 0
                continue
        if item_name == 'HDR_Warnings': #1.警告の後に空のparagraphを挿入
            cmn_para = document.add_paragraph()

        fln_nonewlin = 1

    document.save(word_write_file_name)
    print(targetfile, "変換処理完了")

#------------------------------------------------------------------------
# tkinter GUI build
if __name__ == '__main__':
    root = Tk()
    root.title('PMDA新記載要領XML→docx変換')
    root.geometry("400x150")

    # InputDir Frame
    inpt_frm = ttk.Frame(root, padding=10 )
    inpt_frm.grid()

    # 入力ファイルラベル
    lbl_dir = ttk.Label(inpt_frm, text="入力ファイル")
    lbl_dir.grid(row=0, column=0, sticky=E)
    # 「保存フォルダ」エントリーの作成
    file_name = StringVar()
    infile_box = ttk.Entry(inpt_frm,textvariable=file_name,width=30)
    infile_box.grid(row=0, column=1, sticky=E)
    # 入力ファイル「参照」ボタン
    infile_btn = ttk.Button(inpt_frm,text=u"ファイル参照", command=fileselect_clicked)
    infile_btn.grid(row=0, column=2, sticky=E)
    # 保存先を開くボタン
    btn_opndir = ttk.Button(inpt_frm, text='出力先を開く',command=lambda: OpenFolder(out_dir) )
    btn_opndir.grid(row=1, column=2, sticky=E)

    # SaveDir Frame
#    save_frm = ttk.Frame(root, padding=10 )
#    save_frm.grid(row=1, column=0, sticky=E)
    # 「保存フォルダ」ラベル
#    lbl_dir = ttk.Label(save_frm, text="保存フォルダ", padding=(5, 2))
#    lbl_dir.grid(row=2, column=0, sticky=E)

    # 「保存フォルダ」エントリーの作成
#    sav_dir = StringVar()
#    sav_dir.set(out_dir)
#    IDirEntry = ttk.Entry(save_frm, textvariable=sav_dir, width=30)
#    IDirEntry.grid(row=2, column=1, sticky=E)

    # 「参照」ボタンの作成
#    IDirButton = ttk.Button(save_frm, text="参照", command=dirdialog_clicked)
#    IDirButton.grid(row=2, column=2, sticky=E)
    # 保存先を開くボタン
#    btn_opndir = ttk.Button(
#        save_frm, text='保存先を開く', 
#        command=lambda: OpenFolder(sav_dir.get())
#    )
#    btn_opndir.grid(row=3, column=2)

    # 空行
    blk_frm = ttk.Frame(root, padding=100 ,relief=RAISED)
    blk_frm.grid(row=4, column=0, sticky=E)
    # 「変換を開始」ボタンの作成
    conv_frm = ttk.Frame(root, padding=2 ,relief=RAISED)
    conv_frm.grid(row=5, column=0, sticky=E)
    convButton = ttk.Button(conv_frm, text="XML→docx 変換を開始",command=lambda:convbutton_clicked(file_name.get()),width=50)
    convButton.grid(row=5, column=0, sticky=E)

    root.mainloop()